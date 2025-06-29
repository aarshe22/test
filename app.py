import os
import streamlit as st
from transformers import AutoTokenizer, AutoModelForCausalLM
import torch
from unstructured.partition.auto import partition
import tempfile
import pandas as pd
import PyPDF2
import docx
import GPUtil
import psutil
import time
import json
import glob
import gc
from typing import List, Dict, Any
from pathlib import Path

MODEL_DIR = "./models/mistral-7b"
DATA_DIR = "/data"  # Container path for data folder

# Memory management settings
MAX_CHUNK_SIZE = 4000  # Maximum tokens per chunk
MAX_DOCUMENTS_PER_BATCH = 5  # Maximum documents to process at once
MAX_TOTAL_CHARS = 1000000  # Maximum total characters across all documents

st.set_page_config(page_title="Mistral 7B Chat with Document Upload", layout="wide")

# Initialize session state for conversation memory
if "conversation_history" not in st.session_state:
    st.session_state.conversation_history = []
if "document_context" not in st.session_state:
    st.session_state.document_context = ""
if "context_window" not in st.session_state:
    st.session_state.context_window = 10  # Number of recent messages to keep
if "data_folder_documents" not in st.session_state:
    st.session_state.data_folder_documents = {}
if "document_source" not in st.session_state:
    st.session_state.document_source = "upload"  # "upload" or "data_folder"

@st.cache_resource(show_spinner=True)
def load_model_and_tokenizer():
    tokenizer = AutoTokenizer.from_pretrained(MODEL_DIR)
    model = AutoModelForCausalLM.from_pretrained(
        MODEL_DIR,
        torch_dtype=torch.float16,
        device_map="auto",
        load_in_8bit=True,
    )
    return tokenizer, model

def extract_text_from_file(uploaded_file):
    suffix = uploaded_file.name.split('.')[-1].lower()
    with tempfile.NamedTemporaryFile(delete=False, suffix=f".{suffix}") as tmp:
        tmp.write(uploaded_file.read())
        tmp_path = tmp.name

    if suffix == "pdf":
        elements = partition(tmp_path)
        text = "\n".join(str(el) for el in elements)
    elif suffix in ["doc", "docx"]:
        doc = docx.Document(tmp_path)
        text = "\n".join(p.text for p in doc.paragraphs)
    elif suffix == "csv":
        df = pd.read_csv(tmp_path)
        text = df.to_string()
    else:
        with open(tmp_path, "r", encoding="utf-8") as f:
            text = f.read()

    os.unlink(tmp_path)
    return text

def extract_text_from_file_path(file_path):
    """Extract text from a file given its path."""
    suffix = Path(file_path).suffix.lower()[1:]  # Remove the dot
    
    try:
        if suffix == "pdf":
            elements = partition(file_path)
            text = "\n".join(str(el) for el in elements)
        elif suffix in ["doc", "docx"]:
            doc = docx.Document(file_path)
            text = "\n".join(p.text for p in doc.paragraphs)
        elif suffix == "csv":
            df = pd.read_csv(file_path)
            text = df.to_string()
        elif suffix == "txt":
            with open(file_path, "r", encoding="utf-8") as f:
                text = f.read()
        else:
            # Try to read as text for unknown extensions
            with open(file_path, "r", encoding="utf-8") as f:
                text = f.read()
        
        return text
    except Exception as e:
        return f"Error reading file {file_path}: {str(e)}"

def scan_data_folder():
    """Scan the data folder for supported document types."""
    if not os.path.exists(DATA_DIR):
        return {}
    
    supported_extensions = ['.pdf', '.doc', '.docx', '.csv', '.txt']
    documents = {}
    
    # Recursively scan the data directory
    for root, dirs, files in os.walk(DATA_DIR):
        for file in files:
            file_path = os.path.join(root, file)
            file_ext = Path(file).suffix.lower()
            
            if file_ext in supported_extensions:
                # Get relative path from data directory
                rel_path = os.path.relpath(file_path, DATA_DIR)
                documents[rel_path] = {
                    'path': file_path,
                    'size': os.path.getsize(file_path),
                    'extension': file_ext
                }
    
    return documents

def load_data_folder_documents():
    """Load and extract text from all documents in the data folder with memory management."""
    documents = scan_data_folder()
    extracted_docs = {}
    
    if not documents:
        return extracted_docs
    
    # Sort documents by size to process smaller ones first
    sorted_docs = sorted(documents.items(), key=lambda x: x[1]['size'])
    
    # Limit number of documents processed
    if len(sorted_docs) > MAX_DOCUMENTS_PER_BATCH:
        st.warning(f"Found {len(sorted_docs)} documents. Processing first {MAX_DOCUMENTS_PER_BATCH} documents to prevent memory issues.")
        sorted_docs = sorted_docs[:MAX_DOCUMENTS_PER_BATCH]
    
    progress_bar = st.progress(0)
    status_text = st.empty()
    memory_text = st.empty()
    
    total_chars = 0
    
    for idx, (rel_path, doc_info) in enumerate(sorted_docs):
        status_text.text(f"Processing {rel_path}...")
        
        # Check memory before processing
        allocated, reserved = get_memory_usage()
        memory_text.text(f"GPU Memory: {allocated:.2f}GB allocated, {reserved:.2f}GB reserved")
        
        try:
            text = extract_text_from_file_path(doc_info['path'])
            
            # Check if adding this document would exceed character limit
            if total_chars + len(text) > MAX_TOTAL_CHARS:
                st.warning(f"Document {rel_path} would exceed total character limit. Truncating...")
                text = truncate_text(text, MAX_TOTAL_CHARS - total_chars)
            
            # Chunk the text if it's too large
            chunks = chunk_text(text)
            
            if len(chunks) > 1:
                st.info(f"Document {rel_path} split into {len(chunks)} chunks for memory efficiency.")
            
            extracted_docs[rel_path] = {
                'text': text,
                'chunks': chunks,
                'size': doc_info['size'],
                'extension': doc_info['extension'],
                'estimated_tokens': estimate_tokens(text)
            }
            
            total_chars += len(text)
            
            # Clear GPU memory periodically
            if idx % 2 == 0:
                clear_gpu_memory()
                
        except Exception as e:
            extracted_docs[rel_path] = {
                'text': f"Error extracting text: {str(e)}",
                'chunks': [],
                'size': doc_info['size'],
                'extension': doc_info['extension'],
                'error': True
            }
        
        progress_bar.progress((idx + 1) / len(sorted_docs))
    
    progress_bar.empty()
    status_text.empty()
    memory_text.empty()
    
    # Final memory cleanup
    clear_gpu_memory()
    
    return extracted_docs

def build_conversation_prompt(user_input: str, conversation_history: List[Dict], document_context: str = "") -> str:
    """Build a prompt that includes conversation history and document context."""
    
    # Start with system instruction
    prompt = "You are a helpful AI assistant. Answer questions based on the provided documents and conversation context.\n\n"
    
    # Add document context if available
    if document_context:
        prompt += f"### DOCUMENTS ###\n{document_context}\n\n"
    
    # Add conversation history (limited by context window)
    if conversation_history:
        prompt += "### CONVERSATION HISTORY ###\n"
        for msg in conversation_history[-st.session_state.context_window:]:
            role = "User" if msg["role"] == "user" else "Assistant"
            prompt += f"{role}: {msg['content']}\n"
        prompt += "\n"
    
    # Add current user input
    prompt += f"### CURRENT QUESTION ###\n{user_input}\n\n### ANSWER ###\n"
    
    return prompt

def stream_generate(prompt: str, tokenizer, model, max_new_tokens=500, temperature=0.7):
    """Generate response with real-time token streaming."""
    inputs = tokenizer(prompt, return_tensors="pt").to(model.device)
    
    # Create a placeholder for streaming
    response_placeholder = st.empty()
    full_response = ""
    
    with torch.no_grad():
        # Generate tokens one by one for streaming effect
        generated_ids = model.generate(
            **inputs,
            max_new_tokens=max_new_tokens,
            do_sample=True,
            temperature=temperature,
            pad_token_id=tokenizer.eos_token_id,
            return_dict_in_generate=True,
            output_scores=False,
        )
        
        # Stream the response token by token
        for i in range(len(inputs.input_ids[0]), len(generated_ids.sequences[0])):
            # Decode up to current token
            current_ids = generated_ids.sequences[0][:i+1]
            current_text = tokenizer.decode(current_ids, skip_special_tokens=True)
            
            # Extract only the new response part
            if "### ANSWER ###" in current_text:
                response = current_text.split("### ANSWER ###")[-1].strip()
            else:
                response = current_text.strip()
            
            # Update the placeholder with current response
            response_placeholder.markdown(f"**🤖 Assistant:** {response}")
            
            # Small delay for streaming effect
            time.sleep(0.01)
            
            full_response = response
    
    return full_response

def add_message_to_history(role: str, content: str):
    """Add a message to the conversation history."""
    st.session_state.conversation_history.append({
        "role": role,
        "content": content,
        "timestamp": time.time()
    })
    
    # Trim history if it exceeds context window
    if len(st.session_state.conversation_history) > st.session_state.context_window * 2:
        st.session_state.conversation_history = st.session_state.conversation_history[-st.session_state.context_window:]

def clear_conversation():
    """Clear the conversation history."""
    st.session_state.conversation_history = []
    st.rerun()

def export_conversation():
    """Export conversation history to JSON."""
    if st.session_state.conversation_history:
        timestamp = time.strftime("%Y%m%d_%H%M%S")
        filename = f"conversation_{timestamp}.json"
        
        export_data = {
            "timestamp": timestamp,
            "document_context": st.session_state.document_context,
            "conversation": st.session_state.conversation_history
        }
        
        return json.dumps(export_data, indent=2), filename
    return None, None

def get_system_stats():
    gpu_load = gpu_mem_used = gpu_mem_total = 0
    gpus = GPUtil.getGPUs()
    if gpus:
        gpu = gpus[0]
        gpu_load = round(gpu.load * 100, 1)
        gpu_mem_used = round(gpu.memoryUsed, 1)
        gpu_mem_total = round(gpu.memoryTotal, 1)

    cpu_load = psutil.cpu_percent()
    ram = psutil.virtual_memory()
    ram_used = round(ram.used / (1024 ** 3), 2)
    ram_total = round(ram.total / (1024 ** 3), 2)

    return gpu_load, gpu_mem_used, gpu_mem_total, cpu_load, ram_used, ram_total

def get_conversation_summary(conversation_history: List[Dict]) -> str:
    """Generate a brief summary of the conversation."""
    if not conversation_history:
        return "No conversation yet."
    
    user_messages = [msg["content"] for msg in conversation_history if msg["role"] == "user"]
    assistant_messages = [msg["content"] for msg in conversation_history if msg["role"] == "assistant"]
    
    summary = f"Conversation has {len(conversation_history)} messages:\n"
    summary += f"• {len(user_messages)} user questions\n"
    summary += f"• {len(assistant_messages)} assistant responses\n"
    
    if user_messages:
        summary += f"• Last question: {user_messages[-1][:100]}{'...' if len(user_messages[-1]) > 100 else ''}"
    
    return summary

def chunk_text(text: str, max_chunk_size: int = MAX_CHUNK_SIZE) -> List[str]:
    """Split text into chunks based on token count."""
    if not text or len(text) < max_chunk_size:
        return [text]
    
    # Simple chunking by sentences and paragraphs
    chunks = []
    current_chunk = ""
    
    # Split by paragraphs first
    paragraphs = text.split('\n\n')
    
    for paragraph in paragraphs:
        # If adding this paragraph would exceed chunk size, save current chunk
        if len(current_chunk) + len(paragraph) > max_chunk_size and current_chunk:
            chunks.append(current_chunk.strip())
            current_chunk = ""
        
        # If paragraph itself is too large, split by sentences
        if len(paragraph) > max_chunk_size:
            sentences = paragraph.split('. ')
            for sentence in sentences:
                if len(current_chunk) + len(sentence) > max_chunk_size and current_chunk:
                    chunks.append(current_chunk.strip())
                    current_chunk = ""
                current_chunk += sentence + ". "
        else:
            current_chunk += paragraph + "\n\n"
    
    # Add the last chunk if it has content
    if current_chunk.strip():
        chunks.append(current_chunk.strip())
    
    return chunks

def estimate_tokens(text: str) -> int:
    """Estimate token count for text (rough approximation)."""
    # Rough estimation: 1 token ≈ 4 characters for English text
    return len(text) // 4

def truncate_text(text: str, max_chars: int = MAX_TOTAL_CHARS) -> str:
    """Truncate text if it exceeds maximum character limit."""
    if len(text) <= max_chars:
        return text
    
    # Try to truncate at a sentence boundary
    truncated = text[:max_chars]
    last_period = truncated.rfind('.')
    last_newline = truncated.rfind('\n')
    
    if last_period > last_newline and last_period > max_chars * 0.8:
        return truncated[:last_period + 1]
    elif last_newline > max_chars * 0.8:
        return truncated[:last_newline]
    else:
        return truncated + "\n\n[Text truncated due to size limits]"

def clear_gpu_memory():
    """Clear GPU memory and run garbage collection."""
    if torch.cuda.is_available():
        torch.cuda.empty_cache()
        torch.cuda.synchronize()
    gc.collect()

def get_memory_usage():
    """Get current memory usage information."""
    if torch.cuda.is_available():
        allocated = torch.cuda.memory_allocated() / (1024**3)  # GB
        reserved = torch.cuda.memory_reserved() / (1024**3)   # GB
        return allocated, reserved
    return 0, 0

def main():
    st.title("💡 Mistral 7B Chat with Document Upload")
    
    # Sidebar for controls
    with st.sidebar:
        st.header("⚙️ Settings")
        
        # Document source selection
        st.header("📁 Document Source")
        document_source = st.radio(
            "Choose document source:",
            ["📤 Upload Files", "📂 Data Folder"],
            index=0 if st.session_state.document_source == "upload" else 1,
            help="Upload files manually or use documents from ./data folder"
        )
        st.session_state.document_source = "upload" if document_source == "📤 Upload Files" else "data_folder"
        
        # Context window control
        context_window = st.slider(
            "Conversation Memory (messages)", 
            min_value=5, 
            max_value=20, 
            value=st.session_state.context_window,
            help="Number of recent messages to include in context"
        )
        st.session_state.context_window = context_window
        
        # Model parameters
        st.header("🎛️ Model Parameters")
        temperature = st.slider(
            "Temperature", 
            min_value=0.1, 
            max_value=1.5, 
            value=0.7, 
            step=0.1,
            help="Higher values make output more creative, lower values more focused"
        )
        max_tokens = st.slider(
            "Max Tokens", 
            min_value=100, 
            max_value=1000, 
            value=500, 
            step=50,
            help="Maximum number of tokens to generate"
        )
        
        # Memory management settings
        st.header("💾 Memory Management")
        with st.expander("⚙️ Advanced Memory Settings"):
            st.caption("Adjust these settings if you experience memory issues")
            
            chunk_size = st.slider(
                "Max Chunk Size (chars)", 
                min_value=1000, 
                max_value=8000, 
                value=MAX_CHUNK_SIZE, 
                step=500,
                help="Maximum characters per document chunk"
            )
            
            max_docs = st.slider(
                "Max Documents per Batch", 
                min_value=1, 
                max_value=10, 
                value=MAX_DOCUMENTS_PER_BATCH, 
                step=1,
                help="Maximum documents to process at once"
            )
            
            max_chars = st.slider(
                "Max Total Characters", 
                min_value=500000, 
                max_value=2000000, 
                value=MAX_TOTAL_CHARS, 
                step=100000,
                help="Maximum total characters across all documents"
            )
            
            if st.button("🔄 Clear GPU Memory"):
                clear_gpu_memory()
                st.success("GPU memory cleared!")
        
        # Show current memory usage
        allocated, reserved = get_memory_usage()
        st.metric("GPU Memory Used", f"{allocated:.2f} GB")
        st.metric("GPU Memory Reserved", f"{reserved:.2f} GB")
        
        # Conversation controls
        st.header("💬 Conversation")
        if st.button("🗑️ Clear Conversation"):
            clear_conversation()
        
        # Conversation summary
        if st.session_state.conversation_history:
            with st.expander("📋 Conversation Summary"):
                st.text(get_conversation_summary(st.session_state.conversation_history))
        
        # Export conversation
        export_data, filename = export_conversation()
        if export_data:
            st.download_button(
                label="📥 Export Conversation",
                data=export_data,
                file_name=filename,
                mime="application/json"
            )
        
        # System stats
        st.header("📊 System Stats")
        gpu_load, gpu_mem_used, gpu_mem_total, cpu_load, ram_used, ram_total = get_system_stats()
        st.metric("GPU Utilization", f"{gpu_load}%")
        st.metric("GPU VRAM", f"{gpu_mem_used}/{gpu_mem_total} GB")
        st.metric("CPU Utilization", f"{cpu_load}%")
        st.metric("RAM Usage", f"{ram_used}/{ram_total} GB")

    tokenizer, model = load_model_and_tokenizer()

    # Document handling based on source selection
    if st.session_state.document_source == "upload":
        # Original file upload functionality
        uploaded_files = st.file_uploader("Upload PDF, DOC, DOCX, CSV or TXT files", type=["pdf", "doc", "docx", "csv", "txt"], accept_multiple_files=True)
        
        if uploaded_files:
            # Limit number of uploaded files
            if len(uploaded_files) > MAX_DOCUMENTS_PER_BATCH:
                st.warning(f"Too many files uploaded ({len(uploaded_files)}). Processing first {MAX_DOCUMENTS_PER_BATCH} files to prevent memory issues.")
                uploaded_files = uploaded_files[:MAX_DOCUMENTS_PER_BATCH]
            
            st.info(f"Extracting text from {len(uploaded_files)} files...")
            progress_bar = st.progress(0)
            memory_text = st.empty()
            combined_text = ""
            total_chars = 0
            
            for idx, file in enumerate(uploaded_files):
                # Check memory before processing
                allocated, reserved = get_memory_usage()
                memory_text.text(f"GPU Memory: {allocated:.2f}GB allocated, {reserved:.2f}GB reserved")
                
                text = extract_text_from_file(file)
                
                # Check if adding this document would exceed character limit
                if total_chars + len(text) > MAX_TOTAL_CHARS:
                    st.warning(f"File {file.name} would exceed total character limit. Truncating...")
                    text = truncate_text(text, MAX_TOTAL_CHARS - total_chars)
                
                # Chunk the text if it's too large
                chunks = chunk_text(text)
                if len(chunks) > 1:
                    st.info(f"File {file.name} split into {len(chunks)} chunks for memory efficiency.")
                
                combined_text += f"\n--- Start of {file.name} ---\n{text}\n--- End of {file.name} ---\n"
                total_chars += len(text)
                
                # Clear GPU memory periodically
                if idx % 2 == 0:
                    clear_gpu_memory()
                
                progress_bar.progress((idx + 1) / len(uploaded_files))
            
            progress_bar.empty()
            memory_text.empty()
            
            # Final memory cleanup
            clear_gpu_memory()
            
            st.success("Text extraction complete.")
            
            # Update document context in session state
            st.session_state.document_context = combined_text
            
            # Display document text in expander
            with st.expander("📄 View Document Text"):
                st.text_area("Combined Document Text", combined_text, height=300, key="document_text_display")
    
    else:
        # Data folder functionality
        st.header("📂 Data Folder Documents")
        
        # Check if data folder exists
        if not os.path.exists(DATA_DIR):
            st.warning(f"Data folder not found at {DATA_DIR}. Please ensure the ./data folder is mounted correctly.")
        else:
            # Scan for documents
            documents = scan_data_folder()
            
            if not documents:
                st.info("No supported documents found in the data folder. Supported formats: PDF, DOC, DOCX, CSV, TXT")
            else:
                st.success(f"Found {len(documents)} documents in the data folder.")
                
                # Show document list
                with st.expander(f"📋 Document List ({len(documents)} files)"):
                    for rel_path, doc_info in documents.items():
                        size_mb = doc_info['size'] / (1024 * 1024)
                        st.text(f"📄 {rel_path} ({size_mb:.2f} MB)")
                
                # Load documents button
                if st.button("🔄 Load Documents from Data Folder"):
                    with st.spinner("Loading documents from data folder..."):
                        extracted_docs = load_data_folder_documents()
                        st.session_state.data_folder_documents = extracted_docs
                        
                        # Combine all extracted text
                        combined_text = ""
                        for rel_path, doc_data in extracted_docs.items():
                            if not doc_data.get('error', False):
                                combined_text += f"\n--- Start of {rel_path} ---\n{doc_data['text']}\n--- End of {rel_path} ---\n"
                        
                        st.session_state.document_context = combined_text
                        st.success(f"Loaded {len(extracted_docs)} documents successfully!")
                
                # Show loaded documents
                if st.session_state.data_folder_documents:
                    st.subheader("📄 Loaded Documents")
                    with st.expander("📋 View Loaded Documents"):
                        for rel_path, doc_data in st.session_state.data_folder_documents.items():
                            if doc_data.get('error', False):
                                st.error(f"❌ {rel_path}: {doc_data['text']}")
                            else:
                                chunks_info = f" ({len(doc_data['chunks'])} chunks)" if len(doc_data['chunks']) > 1 else ""
                                tokens_info = f" (~{doc_data['estimated_tokens']} tokens)" if 'estimated_tokens' in doc_data else ""
                                st.success(f"✅ {rel_path}{chunks_info}{tokens_info} ({len(doc_data['text'])} characters)")
                    
                    # Show combined text
                    if st.session_state.document_context:
                        with st.expander("📄 View Combined Document Text"):
                            st.text_area("Combined Document Text", st.session_state.document_context, height=300, key="data_folder_text_display")
                        
                        # Show memory usage
                        allocated, reserved = get_memory_usage()
                        st.info(f"📊 Current GPU Memory: {allocated:.2f}GB allocated, {reserved:.2f}GB reserved")

    # Chat interface
    st.header("💬 Chat Interface")
    
    # Display conversation history in a more compact format
    if st.session_state.conversation_history:
        st.subheader(f"📝 Conversation History ({len(st.session_state.conversation_history)} messages)")
        
        # Create tabs for different views
        tab1, tab2 = st.tabs(["💬 Messages", "📊 Summary"])
        
        with tab1:
            for i, msg in enumerate(st.session_state.conversation_history):
                if msg["role"] == "user":
                    st.markdown(f"**👤 You:** {msg['content']}")
                else:
                    st.markdown(f"**🤖 Assistant:** {msg['content']}")
                st.divider()
        
        with tab2:
            st.text(get_conversation_summary(st.session_state.conversation_history))

    # User input
    user_input = st.text_input("💬 Enter your question or prompt:", key="user_input")
    
    col1, col2, col3 = st.columns([1, 2, 2])
    with col1:
        generate_button = st.button("🚀 Generate Response")
    with col2:
        if st.session_state.conversation_history:
            st.caption(f"📊 Using last {len(st.session_state.conversation_history)} messages as context")
    with col3:
        if st.session_state.document_context:
            st.caption("📄 Document context available")

    if generate_button and user_input:
        # Add user message to history
        add_message_to_history("user", user_input)
        
        # Build prompt with conversation history
        full_prompt = build_conversation_prompt(
            user_input, 
            st.session_state.conversation_history, 
            st.session_state.document_context
        )
        
        with st.spinner("Generating response..."):
            # Pass model parameters to the generation function
            response = stream_generate(full_prompt, tokenizer, model, max_new_tokens=max_tokens, temperature=temperature)
            
            # Add assistant response to history
            add_message_to_history("assistant", response)
            
            # Rerun to update the conversation display
            st.rerun()

if __name__ == "__main__":
    main()

